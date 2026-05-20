"""Generate KOREA_DIGEST_OVERVIEW.docx from the markdown content."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xC9, 0xA9, 0x6E)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ═══════════════════════════════════════════════════════════════════════════
# COVER / TITLE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_paragraph()  # spacer
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CSIS Korea Daily Brief')
run.font.size = Pt(28)
run.font.color.rgb = NAVY
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('System Overview')
run.font.size = Pt(18)
run.font.color.rgb = GRAY

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Automated daily intelligence product for the CSIS Korea Chair')
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.italic = True

doc.add_paragraph()

details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details.add_run('9 Sections  ·  138+ Sources  ·  4 Tiers  ·  ~$12/month')
run.font.size = Pt(11)
run.font.color.rgb = GOLD
run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1: WHAT IT DOES
# ═══════════════════════════════════════════════════════════════════════════

add_heading('1. What It Does', level=1)

add_body(
    'Every morning at 6:00 AM ET, the system automatically produces a '
    'Presidential Daily Brief-style newsletter for senior policymakers, '
    'Korea scholars, and elite journalists. The entire pipeline runs '
    'unattended on GitHub Actions in 3-5 minutes.'
)

steps = [
    ('Collects ', 'articles from 138+ sources across 4 tiers (Korean, English, government, KCNA)'),
    ('Fetches ', 'live market data (KRW/USD, KOSPI, BOK rate, 10Y yield) and polling numbers'),
    ('Loads ', 'CSIS databases (NK-Russia timeline, NK provocations since 1958) for historical context'),
    ('Synthesizes ', 'everything through Claude Opus into a structured 1,200-1,400 word briefing'),
    ('Validates ', 'the output (URL checks, deduplication, word count, hallucination guards)'),
    ('Renders ', 'a styled HTML email optimized for Gmail/Outlook'),
    ('Sends ', 'via Gmail SMTP and publishes to GitHub Pages'),
]
for bold_part, rest in steps:
    add_bullet(rest, bold_prefix=bold_part)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2: THE 9 SECTIONS
# ═══════════════════════════════════════════════════════════════════════════

add_heading('2. The 9 Sections', level=1)

add_table(
    ['#', 'Section', 'What It Contains'],
    [
        ['1', 'Today at a Glance', '3-bullet morning memo — the top-line takeaway'],
        ['2', 'Top Stories + Overnight Flash', '2-4 lead stories + up to 6 overnight items with full sourcing'],
        ['3', 'DPRK Official Statements', 'Kim Jong Un appearance status, official quotes (up to 4 with speaker attribution), senior official activity'],
        ['4', 'Satellite & Location Watch', '11 monitored facilities (Yongbyon, Punggye-ri, Sohae, Sinpo, etc.) with status badges + imagery reports'],
        ['5', 'ROK Government', 'Ministry-by-ministry action cards (Blue House, MOFA, MND, NIS, FSC) + National Assembly + personnel changes'],
        ['6', 'Election Tracker', 'Local election race tracker with party standings and countdown'],
        ['7', 'Business & Economy', 'Corporate news, trade data, US-Korea investment deals'],
        ['8', 'Northeast Asia Watch', 'Japan, China, Russia → Korea developments'],
        ['9', 'Public Sentiment Tracker', 'Presidential approval + party support (Gallup Korea / Realmeter)'],
    ],
    col_widths=[0.4, 2.0, 4.0],
)

add_body('')
add_body(
    'Also included: Market indicators bar (header), Key Stat of the Day, '
    'The Wire (secondary items), Statements & Analysis (op-eds, think tank '
    'pieces, academic papers), On This Day (footer).',
    italic=True,
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3: SOURCE ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════

add_heading('3. Source Architecture — 138+ Feeds', level=1)

# Tier 1
add_heading('Tier 1 — Breaking News (71 feeds)', level=2)
add_body('The raw intelligence intake. Collected in parallel every morning.')

tier1_cats = [
    ('Korean English-language dailies: ', 'Korea Herald, Korea Times, Yonhap English, JoongAng Daily, Chosun English, Hankyoreh English, Dong-A English'),
    ('Korean-language newspapers: ', '조선일보, 한겨레, 동아일보, 경향신문, 뉴스1, 연합뉴스, MBN (Claude translates during analysis)'),
    ('Korean broadcast: ', 'JTBC, KBS, MBC, SBS, YTN, Channel A, Arirang'),
    ('Korean business dailies: ', '매일경제, 한국경제, Korea Economic Daily'),
    ('International correspondents: ', 'WSJ, NYT, WaPo, FT, Reuters, AP, Bloomberg, BBC, CNN, CNBC, Guardian, Al Jazeera'),
    ('Regional Asia: ', 'Nikkei, Japan Times, SCMP, Kyodo, Mainichi, Asahi, CNA'),
    ('US Government: ', 'White House, State Dept, Pentagon, USFK, INDOPACOM, Commerce, Treasury, OFAC, BIS'),
    ('ROK/Japan Government: ', 'ROK MOFA, MOTIE, MND, Japan MOFA'),
    ('US Congress: ', 'Senate Foreign Relations, Senate Armed Services, House Foreign Affairs'),
    ('International orgs: ', 'IAEA, UN Security Council, CISA'),
    ('Reaction layer: ', 'Global Times, Xinhua, TASS, Caixin, China Daily, People\'s Daily'),
]
for bold_part, rest in tier1_cats:
    add_bullet(rest, bold_prefix=bold_part)

# Tier 2
add_heading('Tier 2 — Analysis & Commentary (30 feeds)', level=2)
add_body('Think tanks and policy outlets, ranked by prestige tier.')

add_bullet(
    'CSIS, Brookings, Carnegie, RAND, CFR, 38 North, Foreign Affairs, AEI, KEIA, Beyond Parallel, NK Pro',
    bold_prefix='A-tier: ',
)
add_bullet(
    'Stimson, IISS, ASAN Institute, EAI, Sejong, SIPRI, War on the Rocks, Foreign Policy, Hudson, Heritage, Atlantic Council, NBR, PIIE, USIP, CRS',
    bold_prefix='B-tier: ',
)

# Tier 3
add_heading('Tier 3 — Academic (20 feeds)', level=2)
add_body('Peer-reviewed journals, ranked A+/A/B.')

add_bullet(
    'International Security, International Organization, World Politics, APSR',
    bold_prefix='A+: ',
)
add_bullet(
    'Journal of Conflict Resolution, Security Studies, Asian Survey, Pacific Review, Foreign Affairs, Survival',
    bold_prefix='A: ',
)
add_bullet(
    'Korean Journal of Defense Analysis, North Korean Review, Asian Security, Nonproliferation Review',
    bold_prefix='B: ',
)

# Tier 4
add_heading('Tier 4 — KCNA & State Media (17 feeds)', level=2)
add_body('Direct and indirect KCNA monitoring for official DPRK statements.')

add_bullet('KCNA Watch, kcna.kp, Rodong Sinmun', bold_prefix='Direct: ')
add_bullet('KCNA via Yonhap, Reuters, AP, AFP, BBC, NYT, WaPo', bold_prefix='Wire relays: ')
add_bullet('KCNA via 38 North, Daily NK, NK News, NK Pro', bold_prefix='Specialist: ')

# Additional
add_heading('Additional Data Sources', level=2)
add_bullet('Yahoo Finance (KRW/USD, KOSPI, Brent, S&P 500) + Stooq fallback + BOK ECOS API', bold_prefix='Market data: ')
add_bullet('Google News scraping for Gallup Korea / Realmeter weekly polls + Korean Wikipedia tables', bold_prefix='Polling: ')
add_bullet('NK-Russia cooperation timeline, NK provocations since 1958 (fetched from GitHub repos)', bold_prefix='CSIS databases: ')
add_bullet('5 dedicated feeds monitoring Kim Jong Un appearances via NK Leadership Watch, Daily NK, KCNA Watch', bold_prefix='Kim tracker: ')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4: PIPELINE ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════

add_heading('4. Pipeline Architecture', level=1)

# Flow diagram as styled text
flow = doc.add_paragraph()
flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = flow.add_run('collect.py  →  digest.py  →  run.py  →  render.py  →  send_email.py')
run.font.name = 'Consolas'
run.font.size = Pt(11)
run.font.color.rgb = NAVY
run.bold = True

flow2 = doc.add_paragraph()
flow2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = flow2.add_run('138+ feeds      Claude Opus     Validation     HTML email      Gmail + Pages')
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = GRAY

add_body('')

pipeline_steps = [
    ('collect.py — ',
     'Parallel RSS fetching via ThreadPoolExecutor. Each tier runs concurrently. '
     'Articles are deduplicated, filtered by Korea relevance (regex + keyword matching), '
     'and tagged with source metadata. Market data from Yahoo Finance with Stooq fallback. '
     'Polling scraped from Google News Korean headlines.'),
    ('digest.py — ',
     'Builds a ~20,000-token prompt containing all collected articles, market data, '
     'tracker history (Kim appearances, KCNA rhetoric, facility status), CSIS database context, '
     'and detailed section-by-section instructions. Sends to Claude Opus with prompt caching. '
     'Claude returns structured JSON with all 15+ sections.'),
    ('run.py — ',
     'Orchestrator. Runs validation gate: section counts (hard caps), URL validity '
     '(parallel HEAD checks), duplicate detection (headline similarity + URL dedup), '
     'word count floor (850 minimum), hallucination guards. If critical validation fails, '
     'retries digest generation up to 2x with feedback injected into the prompt.'),
    ('render.py — ',
     'Converts digest JSON to table-based HTML email (~1,400 lines). Every style is inline '
     'for email client compatibility. Responsive via @media queries. Solid background-color '
     'fallbacks for gradient-challenged clients (Gmail, Outlook).'),
    ('send_email.py — ',
     'Gmail SMTP with retry logic. Also writes public/latest.html for GitHub Pages archive.'),
]
for bold_part, rest in pipeline_steps:
    add_bullet(rest, bold_prefix=bold_part)

add_body('')
add_heading('Persistent Trackers', level=2)
add_body('Cached across runs via GitHub Actions:')

add_bullet(
    'Confirmed Kim Jong Un appearances with dates and activities',
    bold_prefix='kim_tracker.json — ',
)
add_bullet(
    'Daily KCNA output history (quotes, watch flags)',
    bold_prefix='kcna_tracker.json — ',
)
add_bullet(
    '11 monitored facility statuses (Yongbyon, Punggye-ri, Sohae, Sinpo, THAAD Seongju, '
    'Tumangang-Khasan, Sinuiju-Dandong, Rason SEZ, Yellow Sea NLL/PMZ, Vostochny/Dunai)',
    bold_prefix='bp_tracker.json — ',
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 5: ANTI-HALLUCINATION
# ═══════════════════════════════════════════════════════════════════════════

add_heading('5. Anti-Hallucination System', level=1)

add_body(
    'The digest serves expert readers (NSC staff, Korea desk officers, senior scholars). '
    'One wrong name or fabricated article destroys credibility. The system enforces 9 guards:'
)

add_table(
    ['Guard', 'What It Does'],
    [
        ['Source-or-Skip', 'Every factual claim must trace to an input article or reference baseline. If neither, it\'s dropped.'],
        ['Think Tank Fabrication Block', 'Hard block on generic-sounding think tank entries. Claude is explicitly told these patterns destroy credibility.'],
        ['URL Validation', 'Every article URL is HEAD-checked in parallel. 404/410 = dead link warning.'],
        ['Duplicate Detection', 'Headline similarity scoring + URL dedup across all sections.'],
        ['Historical Claim Ban', 'Claude cannot cite dates or precedents from training data. Only from today\'s articles or provided reference databases.'],
        ['Arithmetic Lock', 'Pre-calculated totals (market data, facility counts) are passed through verbatim. Claude cannot recalculate.'],
        ['Validation Retry Loop', 'If critical warnings fire, the digest is regenerated up to 2x with validation feedback injected.'],
        ['Word Count Floor', 'Hard minimum 850 words. Blocks truncated outputs.'],
        ['Section Caps', 'Hard max per section (e.g., top_stories: 2-4, overnight: 3-6). Prevents bloat or fabrication padding.'],
    ],
    col_widths=[2.0, 4.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 6: COST ESTIMATE
# ═══════════════════════════════════════════════════════════════════════════

add_heading('6. Cost Estimate', level=1)

add_heading('Per-run cost breakdown', level=2)

add_table(
    ['Component', 'Cost'],
    [
        ['Claude Opus API (primary digest generation)', '~$0.30-0.60 per run'],
        ['Claude Sonnet API (validation retries, if needed)', '~$0.05-0.10 per retry'],
        ['Prompt caching (system prompt cached)', 'Reduces input cost ~50% on retries'],
        ['GitHub Actions', 'Free (public repo)'],
        ['Gmail SMTP', 'Free'],
        ['External APIs (Yahoo Finance, BOK ECOS)', 'Free'],
        ['RSS feeds (Google News, direct)', 'Free'],
    ],
    col_widths=[3.5, 3.0],
)

add_body('')
add_heading('Monthly estimate', level=2)

add_table(
    ['Scenario', 'Monthly Cost'],
    [
        ['Normal (1 run/day, no retries)', '~$9-18/month'],
        ['With retries (~30% of days need 1 retry)', '~$12-22/month'],
        ['Worst case (daily retries + manual re-runs)', '~$25-35/month'],
    ],
    col_widths=[3.5, 3.0],
)

add_body('')
add_body(
    'The only paid dependency is the Anthropic API. Everything else is free infrastructure.',
    bold=True,
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 7: REPLICATION GUIDE
# ═══════════════════════════════════════════════════════════════════════════

add_heading('7. Replication Guide', level=1)

add_heading('Requirements', level=2)

add_table(
    ['Requirement', 'Details'],
    [
        ['GitHub repo', 'Public or private. GitHub Actions runs the pipeline.'],
        ['Anthropic API key', 'Claude Opus access. Set as ANTHROPIC_API_KEY secret.'],
        ['Gmail account', 'With App Password enabled. Set GMAIL_USER, GMAIL_APP_PASS, DIGEST_TO secrets.'],
        ['Python 3.12', 'With anthropic, feedparser, requests, httpx packages.'],
        ['Optional: BOK API key', 'For live BOK economic indicators. Free at ecos.bok.or.kr.'],
        ['Optional: GitHub PAT', 'For database push-back (NK-Russia timeline, provocations).'],
    ],
    col_widths=[2.0, 4.5],
)

add_body('')
add_heading('To replicate for a different region/topic', level=2)

steps = [
    ('1. Fork the repo and modify collect.py: ',
     'Replace TIER1-4_FEEDS with your region\'s sources. Update keyword regex. Adjust market data tickers.'),
    ('2. Modify digest.py: ',
     'Rewrite SYSTEM_PROMPT for your analyst persona and audience. Update section definitions and field specs.'),
    ('3. Modify render.py: ',
     'Update section renderers to match your new digest structure. Adjust branding.'),
    ('4. Modify bp_tracker.json: ',
     'Replace monitored locations with your region\'s facilities/sites.'),
    ('5. Set up GitHub Actions: ',
     'Copy the workflow file. Add secrets to repo settings. Adjust cron schedule for your timezone.'),
]
for bold_part, rest in steps:
    add_bullet(rest, bold_prefix=bold_part)

add_body('')
add_heading('File inventory', level=2)

add_table(
    ['File', 'Lines', 'Role'],
    [
        ['collect.py', '1,653', 'RSS collection, market data, polling scrape'],
        ['render.py', '1,408', 'HTML email renderer'],
        ['run.py', '1,122', 'Orchestrator, validation, postprocessing'],
        ['digest.py', '888', 'Claude API integration, prompt building'],
        ['databases.py', '746', 'CSIS database fetch/push (NK-Russia, provocations)'],
        ['kim_tracker.py', '247', 'Kim Jong Un appearance tracking'],
        ['send_email.py', '210', 'Gmail SMTP + GitHub Pages publish'],
        ['weekly.py', '286', 'Weekly summary generator'],
        ['kcna_tracker.py', '120', 'KCNA rhetoric history'],
        ['bp_tracker.py', '103', 'Facility status persistence'],
        ['Total', '~7,300', ''],
    ],
    col_widths=[1.8, 0.8, 3.8],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 8: Q&A
# ═══════════════════════════════════════════════════════════════════════════

add_heading('8. Q&A', level=1)

qas = [
    ('How accurate are the articles? Can Claude hallucinate stories?',
     'Multiple layers prevent this. Every article must have a real URL from the input feed. '
     'URLs are HEAD-checked. Think tank fabrication is explicitly blocked. If validation catches '
     'a hallucinated entry, the digest is regenerated with feedback. Expert readers would catch a '
     'fabricated CSIS or Brookings piece instantly — the anti-hallucination system is designed for that audience.'),

    ('What happens if the pipeline fails?',
     'The workflow has a fallback cron at 7:00 AM ET. If the primary 6:00 AM run fails, '
     'the fallback fires. If it already succeeded, the fallback detects this and skips. '
     'If both fail, no email is sent — silence is better than a bad product.'),

    ('How does the Korean-language content work?',
     '~20 of the 71 Tier 1 feeds are Korean-language sources (조선일보, 한겨레, JTBC, KBS, etc.). '
     'These articles are collected with their original Korean titles and summaries. Claude Opus '
     'translates and analyzes them during digest generation. This gives the brief coverage that '
     'English-only monitoring misses.'),

    ('What\'s the latency?',
     'Collection takes ~30-60 seconds (parallel RSS fetching). Claude Opus synthesis takes '
     '~60-120 seconds. Validation and rendering take ~15 seconds. Total pipeline: 2-4 minutes '
     'from trigger to inbox.'),

    ('Can readers see past issues?',
     'Yes. Each digest is archived to GitHub Pages. The email includes a "Read online" link. '
     'Daily JSONs are also archived for the weekly summary generator.'),

    ('How is the facility tracker maintained?',
     'bp_tracker.json persists across runs via GitHub Actions cache. When a 38 North or AEI '
     'satellite imagery report appears in the day\'s articles, Claude updates the relevant '
     'facility\'s status and note. Otherwise, the last known status carries forward. The 11 '
     'monitored sites cover nuclear (Yongbyon, Punggye-ri), missile (Sohae, Sinpo), border '
     'crossings (Tumangang-Khasan, Sinuiju-Dandong), economic zones (Rason), military '
     '(THAAD Seongju), maritime (Yellow Sea NLL/PMZ), and logistics (Vostochny/Dunai).'),

    ('Why Claude and not GPT-4 or Gemini?',
     'Claude Opus was chosen for: (1) large context window handling 138+ articles per run, '
     '(2) structured JSON output reliability, (3) prompt caching reducing repeat costs, '
     '(4) strong performance on Korean-language content. The system is model-agnostic in principle — '
     'swap the model ID in digest.py — but the prompt engineering is tuned for Claude\'s strengths.'),

    ('What\'s the difference between this and a human analyst?',
     'This doesn\'t replace an analyst — it replaces the first 2 hours of their morning. '
     'Instead of scanning 70+ sources, pulling market data, checking KCNA, and drafting a '
     'summary, the analyst opens their inbox to a structured brief and spends their time on '
     'analysis, not collection. The system handles breadth; the human provides depth.'),
]

for q, a in qas:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: {q}')
    run.bold = True
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    run = p.add_run(f'A: {a}')
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 9: WHAT'S NEXT
# ═══════════════════════════════════════════════════════════════════════════

add_heading('9. What\'s Next', level=1)

roadmap = [
    ('Gallup Korea direct integration — ',
     'Currently scraping poll numbers from news headlines. Working toward structured polling data extraction.'),
    ('Weekly digest — ',
     'weekly.py aggregates daily JSONs into a weekly summary (architecture exists, refinement ongoing).'),
    ('Database auto-push — ',
     'New NK-Russia cooperation events and provocations are automatically flagged and pushed to CSIS timeline databases.'),
    ('Subscriber management — ',
     'Currently single-recipient. Future: distribution list with subscribe/unsubscribe.'),
]
for bold_part, rest in roadmap:
    add_bullet(rest, bold_prefix=bold_part)

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════

out_path = '/home/user/Daily-Korea-News/KOREA_DIGEST_OVERVIEW.docx'
doc.save(out_path)
print(f'Saved to {out_path}')
